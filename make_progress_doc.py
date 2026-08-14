from pathlib import Path
from docx import Document

root = Path(r"c:\Users\Owner\OneDrive\Desktop\home\学校関係\クラス委員\2年クラス委員")
text_path = root / "04 事前打ち合わせ" / "01_会議持参資料" / "02_当日進行表_案.txt"
doc_path = root / "04 事前打ち合わせ" / "01_会議持参資料" / "02_当日進行表_案.docx"
html_path = root / "04 事前打ち合わせ" / "01_会議持参資料" / "02_当日進行表_案.html"

text = text_path.read_text(encoding="utf-8")

# Word document
word = Document()
word.add_heading("当日進行表（案）", level=1)
word.add_paragraph("開催日：10月3日（土）9時00分")
word.add_paragraph("開催場所：泗水小学校体育館")

table = word.add_table(rows=1, cols=4)
table.style = "Table Grid"
header = table.rows[0].cells
header[0].text = "項目"
header[1].text = "時間"
header[2].text = "内容"
header[3].text = "担当"
rows = [
    ["1. 役員集合", "8:20", "役員全員集合、当日役割の最終確認", ""],
    ["2. 会場準備", "委員到着次第", "準備物や配置などを講師と確認", ""],
    ["3. 受付開始", "随時", "出席確認表（A3）2枚ほどを用意、筆記具を用意、受付用の机を準備", ""],
    ["4. 開会", "9:00", "司会2名を選出、開会の挨拶、講師の先生紹介", ""],
    ["5. レクリエーション実施", "9:10ごろ", "木工ストラップ作り、必要なものは講師確認、約2時間を予定", ""],
    ["6. 閉会", "11:00ごろ", "司会2名が閉会の挨拶、本日の感想を数名に発表、講師へのお礼", ""],
    ["7. お土産配布", "解散時", "お土産の配布とチェック、出席確認表に配布欄を追加", ""],
    ["8. 片付け", "解散後", "会場内の整理・片付け、使用した備品の確認、持ち帰り物のチェック", ""],
]
for row in rows:
    cells = table.add_row().cells
    for i, value in enumerate(row):
        cells[i].text = value

word.add_paragraph("")
word.add_heading("講師へ事前確認したい事項", level=2)
confirm_table = word.add_table(rows=1, cols=2)
confirm_table.style = "Table Grid"
confirm_header = confirm_table.rows[0].cells
confirm_header[0].text = "確認事項"
confirm_header[1].text = "内容"
confirm_rows = [
    ["活動内容", "木工ストラップ作りの流れ、難易度、所要時間の最終確認"],
    ["材料・道具", "必要な材料、道具、消耗品の確認"],
    ["準備物", "1人あたりの準備物や配布量の確認"],
    ["進行タイム", "開始・休憩・終了時刻の目安確認"],
    ["安全対策", "保護者や児童の安全面に関する注意事項"],
    ["写真撮影", "撮影の可否や撮影範囲の確認"],
    ["サポート", "先生側の役割（進行/補助/見守り）"],
    ["緊急時対応", "当日の連絡方法や緊急時対応の確認"],
]
for row in confirm_rows:
    cells = confirm_table.add_row().cells
    for i, value in enumerate(row):
        cells[i].text = value

word.add_paragraph("")
word.add_heading("担当者メモ", level=2)
assign_table = word.add_table(rows=1, cols=2)
assign_table.style = "Table Grid"
assign_header = assign_table.rows[0].cells
assign_header[0].text = "項目"
assign_header[1].text = "担当"
for key in ["受付", "進行", "配布", "写真", "片付け", "先生対応"]:
    row = assign_table.add_row().cells
    row[0].text = key
    row[1].text = ""

word.add_paragraph("")
word.add_heading("備考", level=2)
notes_table = word.add_table(rows=1, cols=2)
notes_table.style = "Table Grid"
notes_header = notes_table.rows[0].cells
notes_header[0].text = "項目"
notes_header[1].text = "内容"
for key in ["必要物品", "連絡事項", "当日の注意点"]:
    row = notes_table.add_row().cells
    row[0].text = key
    row[1].text = ""

word.save(doc_path)

# HTML document
html = '''<!doctype html>
<html lang="ja">
<head>
<meta charset="utf-8">
<title>当日進行表（案）</title>
<style>
body { font-family: "Yu Gothic", Meiryo, sans-serif; margin: 32px; }
.table { border-collapse: collapse; width: 100%; margin-top: 16px; }
th, td { border: 1px solid #333; padding: 8px; vertical-align: top; }
h2 { margin-top: 24px; }
</style>
</head>
<body>
<h1>当日進行表（案）</h1>
<p>開催日：10月3日（土）9時00分<br>開催場所：泗水小学校体育館</p>
<table class="table">
  <tr><th>項目</th><th>時間</th><th>内容</th><th>担当</th></tr>
  <tr><td>1. 役員集合</td><td>8:20</td><td>役員全員集合、当日役割の最終確認</td><td></td></tr>
  <tr><td>2. 会場準備</td><td>委員到着次第</td><td>準備物や配置などを講師と確認</td><td></td></tr>
  <tr><td>3. 受付開始</td><td>随時</td><td>出席確認表（A3）2枚ほどを用意、筆記具を用意、受付用の机を準備</td><td></td></tr>
  <tr><td>4. 開会</td><td>9:00</td><td>司会2名を選出、開会の挨拶、講師の先生紹介</td><td></td></tr>
  <tr><td>5. レクリエーション実施</td><td>9:10ごろ</td><td>木工ストラップ作り、必要なものは講師確認、約2時間を予定</td><td></td></tr>
  <tr><td>6. 閉会</td><td>11:00ごろ</td><td>司会2名が閉会の挨拶、本日の感想を数名に発表、講師へのお礼</td><td></td></tr>
  <tr><td>7. お土産配布</td><td>解散時</td><td>お土産の配布とチェック、出席確認表に配布欄を追加</td><td></td></tr>
  <tr><td>8. 片付け</td><td>解散後</td><td>会場内の整理・片付け、使用した備品の確認、持ち帰り物のチェック</td><td></td></tr>
</table>

<h2>講師へ事前確認したい事項</h2>
<table class="table">
  <tr><th>確認事項</th><th>内容</th></tr>
  <tr><td>活動内容</td><td>木工ストラップ作りの流れ、難易度、所要時間の最終確認</td></tr>
  <tr><td>材料・道具</td><td>必要な材料、道具、消耗品の確認</td></tr>
  <tr><td>準備物</td><td>1人あたりの準備物や配布量の確認</td></tr>
  <tr><td>進行タイム</td><td>開始・休憩・終了時刻の目安確認</td></tr>
  <tr><td>安全対策</td><td>保護者や児童の安全面に関する注意事項</td></tr>
  <tr><td>写真撮影</td><td>撮影の可否や撮影範囲の確認</td></tr>
  <tr><td>サポート</td><td>先生側の役割（進行/補助/見守り）</td></tr>
  <tr><td>緊急時対応</td><td>当日の連絡方法や緊急時対応の確認</td></tr>
</table>

<h2>担当者メモ</h2>
<table class="table">
  <tr><th>項目</th><th>担当</th></tr>
  <tr><td>受付</td><td></td></tr>
  <tr><td>進行</td><td></td></tr>
  <tr><td>配布</td><td></td></tr>
  <tr><td>写真</td><td></td></tr>
  <tr><td>片付け</td><td></td></tr>
  <tr><td>先生対応</td><td></td></tr>
</table>

<h2>備考</h2>
<table class="table">
  <tr><th>項目</th><th>内容</th></tr>
  <tr><td>必要物品</td><td></td></tr>
  <tr><td>連絡事項</td><td></td></tr>
  <tr><td>当日の注意点</td><td></td></tr>
</table>
</body>
</html>
'''
html_path.write_text(html, encoding="utf-8")
print(f"Created: {doc_path}")
print(f"Created: {html_path}")
